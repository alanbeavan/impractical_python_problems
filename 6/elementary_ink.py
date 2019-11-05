#!/usr/bin/env python3.6
"""Hide a message in a word document."""

import docx
from docx.shared import RGBColor, Pt
import my_module as mod

def load_assets(blank_filename, fake_filename, real_filename):
    """Get all the word doc bits in docx formats."""
    blank_doc = docx.Document(blank_filename)   
   
    fake_message = docx.Document(fake_filename)
    fake_paragraphs = []
    for paragraph in fake_message.paragraphs:
        fake_paragraphs.append(paragraph.text)

    real_message = docx.Document(real_filename)
    real_paragraphs = []
    for paragraph in real_message.paragraphs:
        real_paragraphs.append(paragraph.text)

    return blank_doc, fake_paragraphs, real_paragraphs

def add_letterhead(doc):
    """Add the initial bits and bobs to the document."""
    doc.add_heading("Morland Holmes", 0)
    subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
    subtitle.alignment = 1
    doc.add_heading("", 1)
    doc.add_paragraph("December 17, 2015")
    doc.add_paragraph("")

def set_spacing(paragraph):
    """Set line spacing betwween paragraphs (so we can hide the message)."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

def interleave_message(doc, fake, real):
    """Hide the real message in the fake message.

    Write the fake message in the doc where there are lines.
    Write the real message where the fake message is just newlines, in white.
    Save the document.
    """
    length_real = len(real)
    count_real = 0

    for line in fake:
        if count_real < length_real and line == "":
            paragraph = doc.add_paragraph(real[count_real])
            paragraph_index = len(doc.paragraphs) - 1
            run = doc.paragraphs[paragraph_index].runs[0]
            font = run.font
            font.color.rgb = RGBColor(255, 255, 255) #Currently red
            count_real += 1
        else:
            paragraph = doc.add_paragraph(line)
        set_spacing(paragraph)

    doc.save("ciphertext.docx")


def main():
    """Do the things."""
    blank_doc, fake_list, real_list = load_assets("template.docx",
                                                  "fakeMessage.docx",
                                                  "realMessage_Vig.docx")
    add_letterhead(blank_doc)
    interleave_message(blank_doc, fake_list, real_list)
    print("Done!")
    
if __name__ == "__main__":
    main()
